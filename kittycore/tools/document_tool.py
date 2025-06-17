"""
DocumentTool для KittyCore 3.0 - Универсальный инструмент для работы с документами
Поддерживает PDF, DOCX, TXT, RTF, ODT, CSV, JSON, XML и другие форматы
Использует многоуровневую архитектуру с fallback-стратегиями
"""

import os
import io
import json
import csv
import tempfile
import logging
import mimetypes
from pathlib import Path
from typing import Dict, List, Any, Optional, Union, Tuple, BinaryIO
from dataclasses import dataclass, field
from enum import Enum
import asyncio
from datetime import datetime

# Стандартные библиотеки для обработки документов
import base64
import hashlib
import zipfile
import tarfile

# Базовый класс инструмента
from .base_tool import BaseTool

# Настройка логирования
logger = logging.getLogger(__name__)


class DocumentFormat(Enum):
    """Поддерживаемые форматы документов"""
    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    TXT = "txt"
    RTF = "rtf"
    ODT = "odt"
    CSV = "csv"
    JSON = "json"
    XML = "xml"
    HTML = "html"
    MARKDOWN = "md"
    EXCEL = "xlsx"
    POWERPOINT = "pptx"
    IMAGE = "image"  # jpg, png, etc.
    UNKNOWN = "unknown"


class ProcessingStrategy(Enum):
    """Стратегии обработки документов"""
    AUTO = "auto"           # Автоматический выбор лучшей стратегии
    LIGHTWEIGHT = "lightweight"  # Быстрая обработка, базовые библиотеки
    ADVANCED = "advanced"   # Продвинутая обработка, сложные алгоритмы
    CLOUD = "cloud"        # Облачные сервисы (AWS, Google, Azure)
    HYBRID = "hybrid"      # Комбинация локальных и облачных методов


@dataclass
class DocumentMetadata:
    """Метаданные документа"""
    filename: str
    format: DocumentFormat
    size_bytes: int
    mime_type: str
    encoding: Optional[str] = None
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    character_count: Optional[int] = None
    created_date: Optional[datetime] = None
    modified_date: Optional[datetime] = None
    author: Optional[str] = None
    title: Optional[str] = None
    subject: Optional[str] = None
    checksum: Optional[str] = None
    processing_time: Optional[float] = None
    confidence_score: Optional[float] = None
    errors: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)


@dataclass 
class ExtractionResult:
    """Результат извлечения данных из документа"""
    text: str
    metadata: DocumentMetadata
    tables: List[Dict[str, Any]] = field(default_factory=list)
    images: List[Dict[str, Any]] = field(default_factory=list)
    structured_data: Dict[str, Any] = field(default_factory=dict)
    raw_data: Optional[bytes] = None
    success: bool = True
    extraction_method: str = "unknown"
    processing_details: Dict[str, Any] = field(default_factory=dict)


class DocumentProcessor:
    """Базовый класс для обработчиков документов"""
    
    def __init__(self):
        self.name = self.__class__.__name__
        self.supported_formats: List[DocumentFormat] = []
        
    def can_process(self, file_format: DocumentFormat) -> bool:
        """Проверяет, может ли процессор обработать данный формат"""
        return file_format in self.supported_formats
        
    async def process(self, file_data: bytes, metadata: DocumentMetadata) -> ExtractionResult:
        """Асинхронная обработка документа"""
        raise NotImplementedError("Subclasses must implement process method")
        
    def validate_file(self, file_data: bytes, metadata: DocumentMetadata) -> Tuple[bool, str]:
        """Валидация файла перед обработкой"""
        if not file_data:
            return False, "Файл пустой"
        if metadata.size_bytes > 100 * 1024 * 1024:  # 100MB
            return False, "Файл слишком большой (>100MB)"
        return True, "OK"


class DocumentFormatDetector:
    """Умный детектор форматов документов"""
    
    # Magic bytes для определения формата по содержимому
    MAGIC_BYTES = {
        b'%PDF': DocumentFormat.PDF,
        b'PK\x03\x04': DocumentFormat.DOCX,  # ZIP-based formats
        b'\xd0\xcf\x11\xe0': DocumentFormat.DOC,  # MS Office old format
        b'{\rtf': DocumentFormat.RTF,
        b'PK': DocumentFormat.ODT,  # OpenDocument также ZIP-based
        b'\x89PNG': DocumentFormat.IMAGE,
        b'\xff\xd8\xff': DocumentFormat.IMAGE,  # JPEG
        b'GIF8': DocumentFormat.IMAGE,
        b'<html': DocumentFormat.HTML,
        b'<!DOCTYPE': DocumentFormat.HTML,
        b'<?xml': DocumentFormat.XML,
    }
    
    # MIME типы для дополнительной проверки
    MIME_MAPPING = {
        'application/pdf': DocumentFormat.PDF,
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': DocumentFormat.DOCX,
        'application/msword': DocumentFormat.DOC,
        'application/rtf': DocumentFormat.RTF,
        'application/vnd.oasis.opendocument.text': DocumentFormat.ODT,
        'text/plain': DocumentFormat.TXT,
        'text/csv': DocumentFormat.CSV,
        'application/json': DocumentFormat.JSON,
        'application/xml': DocumentFormat.XML,
        'text/xml': DocumentFormat.XML,
        'text/html': DocumentFormat.HTML,
        'text/markdown': DocumentFormat.MARKDOWN,
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': DocumentFormat.EXCEL,
        'application/vnd.openxmlformats-officedocument.presentationml.presentation': DocumentFormat.POWERPOINT,
        'image/jpeg': DocumentFormat.IMAGE,
        'image/png': DocumentFormat.IMAGE,
        'image/gif': DocumentFormat.IMAGE,
        'image/bmp': DocumentFormat.IMAGE,
        'image/tiff': DocumentFormat.IMAGE,
    }
    
    @classmethod
    def detect_format(cls, file_data: bytes, filename: str = None, mime_type: str = None) -> DocumentFormat:
        """Определяет формат документа по содержимому, имени файла и MIME типу"""
        
        # 1. Проверка по magic bytes (наиболее надёжно)
        for magic, doc_format in cls.MAGIC_BYTES.items():
            if file_data.startswith(magic):
                # Дополнительная проверка для ZIP-based форматов
                if magic == b'PK\x03\x04' and filename:
                    return cls._detect_zip_based_format(file_data, filename)
                return doc_format
        
        # 2. Проверка по MIME типу
        if mime_type and mime_type in cls.MIME_MAPPING:
            return cls.MIME_MAPPING[mime_type]
        
        # 3. Проверка по расширению файла
        if filename:
            extension = Path(filename).suffix.lower().lstrip('.')
            for doc_format in DocumentFormat:
                if doc_format.value == extension:
                    return doc_format
        
        # 4. Эвристическая проверка содержимого
        return cls._heuristic_detection(file_data)
    
    @classmethod
    def _detect_zip_based_format(cls, file_data: bytes, filename: str) -> DocumentFormat:
        """Определяет конкретный формат для ZIP-based файлов"""
        try:
            with io.BytesIO(file_data) as bio:
                with zipfile.ZipFile(bio, 'r') as zf:
                    filenames = zf.namelist()
                    
                    # Проверка на DOCX
                    if 'word/document.xml' in filenames:
                        return DocumentFormat.DOCX
                    
                    # Проверка на XLSX
                    if 'xl/workbook.xml' in filenames:
                        return DocumentFormat.EXCEL
                    
                    # Проверка на PPTX
                    if 'ppt/presentation.xml' in filenames:
                        return DocumentFormat.POWERPOINT
                    
                    # Проверка на ODT
                    if 'content.xml' in filenames and 'META-INF/manifest.xml' in filenames:
                        return DocumentFormat.ODT
                        
        except (zipfile.BadZipFile, Exception):
            pass
        
        # Fallback по расширению
        ext = Path(filename).suffix.lower()
        if ext == '.docx':
            return DocumentFormat.DOCX
        elif ext == '.xlsx':
            return DocumentFormat.EXCEL
        elif ext == '.pptx':
            return DocumentFormat.POWERPOINT
        elif ext == '.odt':
            return DocumentFormat.ODT
            
        return DocumentFormat.UNKNOWN
    
    @classmethod
    def _heuristic_detection(cls, file_data: bytes) -> DocumentFormat:
        """Эвристическое определение формата по содержимому"""
        try:
            # Попытка декодировать как текст
            text = file_data.decode('utf-8', errors='ignore')[:1000].lower()
            
            # JSON detection
            if text.strip().startswith(('{', '[')):
                try:
                    json.loads(file_data.decode('utf-8'))
                    return DocumentFormat.JSON
                except:
                    pass
            
            # XML detection
            if text.strip().startswith('<?xml') or '<' in text and '>' in text:
                return DocumentFormat.XML
            
            # HTML detection
            if any(tag in text for tag in ['<html', '<head', '<body', '<!doctype']):
                return DocumentFormat.HTML
            
            # Markdown detection
            if any(marker in text for marker in ['#', '**', '__', '[', '](', '```']):
                return DocumentFormat.MARKDOWN
            
            # CSV detection (простая эвристика)
            lines = text.split('\n')[:10]
            if len(lines) > 1:
                first_line_commas = lines[0].count(',')
                if first_line_commas > 0:
                    # Проверяем консистентность количества запятых
                    consistent = sum(1 for line in lines[1:] if line.count(',') == first_line_commas)
                    if consistent > len(lines) * 0.7:  # 70% строк имеют одинаковое количество запятых
                        return DocumentFormat.CSV
            
            # По умолчанию - текстовый файл
            return DocumentFormat.TXT
            
        except UnicodeDecodeError:
            # Бинарный файл неизвестного формата
            return DocumentFormat.UNKNOWN


class DocumentUtils:
    """Вспомогательные утилиты для работы с документами"""
    
    @staticmethod
    def calculate_checksum(data: bytes) -> str:
        """Вычисляет MD5 хеш для данных"""
        return hashlib.md5(data).hexdigest()
    
    @staticmethod
    def detect_encoding(data: bytes) -> str:
        """Определяет кодировку текстового файла"""
        try:
            # Попытка определить кодировку
            import chardet
            result = chardet.detect(data)
            return result.get('encoding', 'utf-8') if result else 'utf-8'
        except ImportError:
            # Если chardet не установлен, используем простую эвристику
            try:
                data.decode('utf-8')
                return 'utf-8'
            except UnicodeDecodeError:
                try:
                    data.decode('cp1251')
                    return 'cp1251'
                except UnicodeDecodeError:
                    return 'latin1'
    
    @staticmethod
    def sanitize_filename(filename: str) -> str:
        """Очищает имя файла от опасных символов"""
        import re
        # Удаляем опасные символы и ограничиваем длину
        clean_name = re.sub(r'[<>:"/\\|?*]', '_', filename)
        clean_name = clean_name[:255]  # Ограничение длины
        return clean_name
    
    @staticmethod
    def create_temp_file(data: bytes, suffix: str = None) -> str:
        """Создаёт временный файл с данными"""
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
            tmp_file.write(data)
            return tmp_file.name
    
    @staticmethod
    def cleanup_temp_file(filepath: str) -> None:
        """Безопасно удаляет временный файл"""
        try:
            if os.path.exists(filepath):
                os.unlink(filepath)
        except OSError as e:
            logger.warning(f"Failed to cleanup temp file {filepath}: {e}")
    
    @staticmethod
    def validate_file_size(data: bytes, max_size_mb: int = 100) -> Tuple[bool, str]:
        """Проверяет размер файла"""
        size_mb = len(data) / (1024 * 1024)
        if size_mb > max_size_mb:
            return False, f"Файл слишком большой: {size_mb:.1f}MB (максимум {max_size_mb}MB)"
        return True, "OK"
    
    @staticmethod
    def count_words(text: str) -> int:
        """Подсчитывает количество слов в тексте"""
        import re
        words = re.findall(r'\b\w+\b', text)
        return len(words)
    
    @staticmethod
    def extract_metadata_from_content(text: str) -> Dict[str, Any]:
        """Извлекает базовые метаданные из текста"""
        return {
            'character_count': len(text),
            'word_count': DocumentUtils.count_words(text),
            'line_count': text.count('\n') + 1,
            'paragraph_count': len([p for p in text.split('\n\n') if p.strip()]),
        } 


class OCREngine(Enum):
    """Движки OCR для распознавания текста"""
    TESSERACT = "tesseract"          # Open-source OCR
    EASYOCR = "easyocr"             # Deep learning OCR
    PADDLEOCR = "paddleocr"         # PaddlePaddle OCR
    DOCTR = "doctr"                 # Deep learning OCR
    CLOUD_VISION = "cloud_vision"    # Google Cloud Vision
    AWS_TEXTRACT = "aws_textract"    # AWS Textract
    AZURE_VISION = "azure_vision"    # Azure Computer Vision
    MULTIMODAL_LLM = "multimodal_llm" # GPT-4V, Claude 3, etc.


@dataclass
class OCRResult:
    """Результат OCR обработки"""
    text: str
    confidence: float
    engine: OCREngine
    processing_time: float
    word_confidences: List[Tuple[str, float]] = field(default_factory=list)
    bounding_boxes: List[Dict[str, Any]] = field(default_factory=list)
    language: Optional[str] = None
    errors: List[str] = field(default_factory=list)


class BaseOCRProcessor(DocumentProcessor):
    """Базовый класс для OCR процессоров"""
    
    def __init__(self, engine: OCREngine):
        super().__init__()
        self.engine = engine
        self.supported_formats = [DocumentFormat.IMAGE, DocumentFormat.PDF]
        
    async def process(self, file_data: bytes, metadata: DocumentMetadata) -> ExtractionResult:
        """Базовая обработка с OCR"""
        start_time = datetime.now()
        
        try:
            # Выполняем OCR
            ocr_result = await self.perform_ocr(file_data, metadata)
            
            # Создаём результат
            processing_time = (datetime.now() - start_time).total_seconds()
            
            result_metadata = metadata
            result_metadata.processing_time = processing_time
            result_metadata.confidence_score = ocr_result.confidence
            
            return ExtractionResult(
                text=ocr_result.text,
                metadata=result_metadata,
                success=True,
                extraction_method=f"OCR_{self.engine.value}",
                processing_details={
                    "ocr_engine": self.engine.value,
                    "confidence": ocr_result.confidence,
                    "language": ocr_result.language,
                    "word_count": len(ocr_result.word_confidences),
                    "processing_time": processing_time
                }
            )
            
        except Exception as e:
            logger.error(f"OCR processing failed with {self.engine.value}: {e}")
            return ExtractionResult(
                text="",
                metadata=metadata,
                success=False,
                extraction_method=f"OCR_{self.engine.value}_FAILED",
                processing_details={"error": str(e)}
            )
    
    async def perform_ocr(self, file_data: bytes, metadata: DocumentMetadata) -> OCRResult:
        """Выполняет OCR - должен быть переопределён в наследниках"""
        raise NotImplementedError("Subclasses must implement perform_ocr")


class TesseractOCRProcessor(BaseOCRProcessor):
    """OCR процессор на основе Tesseract"""
    
    def __init__(self, language: str = 'rus+eng'):
        super().__init__(OCREngine.TESSERACT)
        self.language = language
        self._tesseract_available = None
        
    def _check_tesseract_availability(self) -> bool:
        """Проверяет доступность Tesseract"""
        if self._tesseract_available is None:
            try:
                import pytesseract
                from PIL import Image
                # Простая проверка
                pytesseract.get_tesseract_version()
                self._tesseract_available = True
            except (ImportError, Exception):
                self._tesseract_available = False
                logger.warning("Tesseract OCR not available. Install: pip install pytesseract pillow")
        return self._tesseract_available
    
    async def perform_ocr(self, file_data: bytes, metadata: DocumentMetadata) -> OCRResult:
        """Выполняет OCR с помощью Tesseract"""
        if not self._check_tesseract_availability():
            raise Exception("Tesseract OCR not available")
            
        import pytesseract
        from PIL import Image
        import time
        
        start_time = time.time()
        
        try:
            # Если это PDF, конвертируем в изображения
            if metadata.format == DocumentFormat.PDF:
                images = self._pdf_to_images(file_data)
            else:
                # Открываем как изображение
                image = Image.open(io.BytesIO(file_data))
                images = [image]
            
            all_text = []
            all_confidences = []
            
            for image in images:
                # Основной OCR
                text = pytesseract.image_to_string(image, lang=self.language)
                all_text.append(text)
                
                # Детальная информация с координатами и уверенностью
                try:
                    data = pytesseract.image_to_data(image, lang=self.language, output_type=pytesseract.Output.DICT)
                    word_confidences = []
                    
                    for i, word in enumerate(data['text']):
                        if word.strip():
                            confidence = float(data['conf'][i])
                            if confidence > 0:  # Tesseract возвращает -1 для нераспознанных
                                word_confidences.append((word, confidence))
                    
                    all_confidences.extend(word_confidences)
                except Exception as e:
                    logger.warning(f"Failed to get detailed OCR data: {e}")
            
            # Объединяем результаты
            combined_text = '\n\n'.join(all_text)
            
            # Вычисляем среднюю уверенность
            if all_confidences:
                avg_confidence = sum(conf for _, conf in all_confidences) / len(all_confidences)
            else:
                avg_confidence = 0.0
            
            processing_time = time.time() - start_time
            
            return OCRResult(
                text=combined_text,
                confidence=avg_confidence / 100.0,  # Нормализуем к 0-1
                engine=self.engine,
                processing_time=processing_time,
                word_confidences=all_confidences,
                language=self.language
            )
            
        except Exception as e:
            raise Exception(f"Tesseract OCR failed: {e}")
    
    def _pdf_to_images(self, pdf_data: bytes) -> List:
        """Конвертирует PDF в изображения"""
        try:
            from pdf2image import convert_from_bytes
            return convert_from_bytes(pdf_data)
        except ImportError:
            logger.warning("pdf2image not available. Install: pip install pdf2image")
            raise Exception("PDF to image conversion not available")


class EasyOCRProcessor(BaseOCRProcessor):
    """OCR процессор на основе EasyOCR"""
    
    def __init__(self, languages: List[str] = None):
        super().__init__(OCREngine.EASYOCR)
        self.languages = languages or ['ru', 'en']
        self._easyocr_reader = None
        
    def _get_reader(self):
        """Ленивая инициализация EasyOCR reader"""
        if self._easyocr_reader is None:
            try:
                import easyocr
                self._easyocr_reader = easyocr.Reader(self.languages)
            except ImportError:
                raise Exception("EasyOCR not available. Install: pip install easyocr")
        return self._easyocr_reader
    
    async def perform_ocr(self, file_data: bytes, metadata: DocumentMetadata) -> OCRResult:
        """Выполняет OCR с помощью EasyOCR"""
        import time
        
        start_time = time.time()
        reader = self._get_reader()
        
        try:
            # EasyOCR может работать с изображениями напрямую
            if metadata.format == DocumentFormat.PDF:
                # Для PDF нужно конвертировать в изображения
                images = self._pdf_to_images(file_data)
                all_results = []
                for image in images:
                    # Конвертируем PIL в numpy array
                    import numpy as np
                    img_array = np.array(image)
                    results = reader.readtext(img_array, detail=1)
                    all_results.extend(results)
            else:
                # Для изображений работаем напрямую
                results = reader.readtext(file_data, detail=1)
                all_results = results
            
            # Извлекаем текст и уверенность
            text_parts = []
            word_confidences = []
            
            for (bbox, text, confidence) in all_results:
                if text.strip():
                    text_parts.append(text)
                    word_confidences.append((text, confidence))
            
            combined_text = ' '.join(text_parts)
            avg_confidence = sum(conf for _, conf in word_confidences) / len(word_confidences) if word_confidences else 0.0
            
            processing_time = time.time() - start_time
            
            return OCRResult(
                text=combined_text,
                confidence=avg_confidence,
                engine=self.engine,
                processing_time=processing_time,
                word_confidences=word_confidences,
                language='+'.join(self.languages)
            )
            
        except Exception as e:
            raise Exception(f"EasyOCR failed: {e}")
    
    def _pdf_to_images(self, pdf_data: bytes) -> List:
        """Конвертирует PDF в изображения"""
        try:
            from pdf2image import convert_from_bytes
            return convert_from_bytes(pdf_data)
        except ImportError:
            raise Exception("PDF to image conversion not available")


class MultimodalLLMOCRProcessor(BaseOCRProcessor):
    """OCR процессор на основе мультимодальных LLM"""
    
    def __init__(self, api_key: str = None, model: str = "gpt-4o"):
        super().__init__(OCREngine.MULTIMODAL_LLM)
        self.api_key = api_key
        self.model = model
        
    async def perform_ocr(self, file_data: bytes, metadata: DocumentMetadata) -> OCRResult:
        """Выполняет OCR с помощью мультимодальной LLM"""
        import time
        
        if not self.api_key:
            raise Exception("API key required for LLM OCR")
        
        start_time = time.time()
        
        try:
            # Кодируем изображение в base64
            if metadata.format == DocumentFormat.PDF:
                # Для PDF берём первую страницу
                images = self._pdf_to_images(file_data)
                if images:
                    image_data = self._image_to_base64(images[0])
                else:
                    raise Exception("Failed to extract image from PDF")
            else:
                image_data = base64.b64encode(file_data).decode('utf-8')
            
            # Вызываем LLM API
            if self.model.startswith('gpt'):
                text = await self._call_openai_vision(image_data)
            elif self.model.startswith('claude'):
                text = await self._call_claude_vision(image_data)
            else:
                raise Exception(f"Unsupported model: {self.model}")
            
            processing_time = time.time() - start_time
            
            return OCRResult(
                text=text,
                confidence=0.9,  # LLM обычно довольно точны
                engine=self.engine,
                processing_time=processing_time,
                language="auto-detected"
            )
            
        except Exception as e:
            raise Exception(f"LLM OCR failed: {e}")
    
    async def _call_openai_vision(self, image_base64: str) -> str:
        """Вызов OpenAI GPT-4V для OCR"""
        try:
            import openai
            
            client = openai.OpenAI(api_key=self.api_key)
            
            response = client.chat.completions.create(
                model=self.model,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": "Extract all text from this image. Preserve formatting and structure. Return only the extracted text without any additional comments."
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/jpeg;base64,{image_base64}"
                                }
                            }
                        ]
                    }
                ],
                max_tokens=4000
            )
            
            return response.choices[0].message.content
            
        except ImportError:
            raise Exception("OpenAI library not available. Install: pip install openai")
    
    async def _call_claude_vision(self, image_base64: str) -> str:
        """Вызов Claude для OCR"""
        try:
            import anthropic
            
            client = anthropic.Anthropic(api_key=self.api_key)
            
            response = client.messages.create(
                model=self.model,
                max_tokens=4000,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image",
                                "source": {
                                    "type": "base64",
                                    "media_type": "image/jpeg",
                                    "data": image_base64
                                }
                            },
                            {
                                "type": "text",
                                "text": "Extract all text from this image. Preserve formatting and structure. Return only the extracted text without any additional comments."
                            }
                        ]
                    }
                ]
            )
            
            return response.content[0].text
            
        except ImportError:
            raise Exception("Anthropic library not available. Install: pip install anthropic")
    
    def _pdf_to_images(self, pdf_data: bytes) -> List:
        """Конвертирует PDF в изображения"""
        try:
            from pdf2image import convert_from_bytes
            return convert_from_bytes(pdf_data, first_page=1, last_page=1)  # Только первая страница
        except ImportError:
            raise Exception("PDF to image conversion not available")
    
    def _image_to_base64(self, image) -> str:
        """Конвертирует PIL изображение в base64"""
        buffer = io.BytesIO()
        image.save(buffer, format='JPEG')
        return base64.b64encode(buffer.getvalue()).decode('utf-8')


class OCROrchestrator:
    """Оркестратор OCR - выбирает лучший движок и комбинирует результаты"""
    
    def __init__(self):
        self.processors: Dict[OCREngine, BaseOCRProcessor] = {}
        self._initialize_processors()
    
    def _initialize_processors(self):
        """Инициализирует доступные OCR процессоры"""
        # Tesseract
        try:
            self.processors[OCREngine.TESSERACT] = TesseractOCRProcessor()
        except Exception:
            logger.info("Tesseract OCR not available")
        
        # EasyOCR
        try:
            self.processors[OCREngine.EASYOCR] = EasyOCRProcessor()
        except Exception:
            logger.info("EasyOCR not available")
    
    def add_llm_processor(self, api_key: str, model: str = "gpt-4o"):
        """Добавляет LLM OCR процессор"""
        self.processors[OCREngine.MULTIMODAL_LLM] = MultimodalLLMOCRProcessor(api_key, model)
    
    async def process_with_best_engine(self, file_data: bytes, metadata: DocumentMetadata) -> ExtractionResult:
        """Обрабатывает документ лучшим доступным OCR движком"""
        if not self.processors:
            raise Exception("No OCR processors available")
        
        # Приоритет движков (можно настраивать)
        priority_order = [
            OCREngine.MULTIMODAL_LLM,  # Наивысший приоритет для LLM
            OCREngine.EASYOCR,         # Хорошая точность
            OCREngine.TESSERACT,       # Fallback
        ]
        
        for engine in priority_order:
            if engine in self.processors:
                try:
                    processor = self.processors[engine]
                    result = await processor.process(file_data, metadata)
                    if result.success and result.text.strip():
                        return result
                except Exception as e:
                    logger.warning(f"OCR engine {engine.value} failed: {e}")
                    continue
        
        # Если все движки не сработали
        raise Exception("All OCR engines failed")
    
    async def process_with_multiple_engines(self, file_data: bytes, metadata: DocumentMetadata) -> List[ExtractionResult]:
        """Обрабатывает документ несколькими OCR движками для сравнения"""
        results = []
        
        for engine, processor in self.processors.items():
            try:
                result = await processor.process(file_data, metadata)
                results.append(result)
            except Exception as e:
                logger.warning(f"OCR engine {engine.value} failed: {e}")
        
        return results 


class PDFProcessor(DocumentProcessor):
    """Процессор для PDF документов с множественными стратегиями"""
    
    def __init__(self):
        super().__init__()
        self.supported_formats = [DocumentFormat.PDF]
        self._libraries_checked = False
        self._available_libraries = {}
        
    def _check_libraries(self):
        """Проверяет доступные библиотеки для работы с PDF"""
        if self._libraries_checked:
            return
            
        # PyMuPDF (fitz)
        try:
            import fitz
            self._available_libraries['fitz'] = True
        except ImportError:
            self._available_libraries['fitz'] = False
            
        # PyPDF2/pypdf
        try:
            import PyPDF2
            self._available_libraries['pypdf2'] = True
        except ImportError:
            try:
                import pypdf
                self._available_libraries['pypdf'] = True
            except ImportError:
                self._available_libraries['pypdf2'] = False
                
        # PDFMiner
        try:
            from pdfminer.high_level import extract_text
            self._available_libraries['pdfminer'] = True
        except ImportError:
            self._available_libraries['pdfminer'] = False
            
        # pdfplumber
        try:
            import pdfplumber
            self._available_libraries['pdfplumber'] = True
        except ImportError:
            self._available_libraries['pdfplumber'] = False
            
        self._libraries_checked = True
    
    async def process(self, file_data: bytes, metadata: DocumentMetadata) -> ExtractionResult:
        """Обрабатывает PDF документ с fallback стратегиями"""
        self._check_libraries()
        
        start_time = datetime.now()
        
        # Приоритет методов извлечения
        methods = [
            ('fitz', self._extract_with_fitz),
            ('pdfplumber', self._extract_with_pdfplumber),
            ('pdfminer', self._extract_with_pdfminer),
            ('pypdf2', self._extract_with_pypdf2),
        ]
        
        for method_name, method_func in methods:
            if self._available_libraries.get(method_name, False):
                try:
                    text, tables, images = await method_func(file_data)
                    
                    if text.strip():  # Если получили текст
                        processing_time = (datetime.now() - start_time).total_seconds()
                        
                        # Обновляем метаданные
                        metadata.processing_time = processing_time
                        metadata.character_count = len(text)
                        metadata.word_count = DocumentUtils.count_words(text)
                        
                        return ExtractionResult(
                            text=text,
                            metadata=metadata,
                            tables=tables,
                            images=images,
                            success=True,
                            extraction_method=f"PDF_{method_name}",
                            processing_details={
                                "method": method_name,
                                "processing_time": processing_time,
                                "page_count": getattr(metadata, 'page_count', None),
                                "table_count": len(tables),
                                "image_count": len(images)
                            }
                        )
                        
                except Exception as e:
                    logger.warning(f"PDF extraction with {method_name} failed: {e}")
                    continue
        
        # Если все методы не сработали, возвращаем ошибку
        return ExtractionResult(
            text="",
            metadata=metadata,
            success=False,
            extraction_method="PDF_ALL_FAILED",
            processing_details={"error": "All PDF extraction methods failed"}
        )
    
    async def _extract_with_fitz(self, file_data: bytes) -> Tuple[str, List[Dict], List[Dict]]:
        """Извлечение с помощью PyMuPDF (fitz)"""
        import fitz
        
        doc = fitz.open(stream=file_data, filetype="pdf")
        text_parts = []
        tables = []
        images = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # Извлекаем текст
            page_text = page.get_text()
            text_parts.append(page_text)
            
            # Извлекаем изображения
            image_list = page.get_images()
            for img_index, img in enumerate(image_list):
                try:
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    images.append({
                        "page": page_num + 1,
                        "index": img_index,
                        "format": base_image["ext"],
                        "size": len(base_image["image"]),
                        "data": base64.b64encode(base_image["image"]).decode()
                    })
                except Exception as e:
                    logger.warning(f"Failed to extract image {img_index} from page {page_num}: {e}")
        
        doc.close()
        return '\n\n'.join(text_parts), tables, images
    
    async def _extract_with_pdfplumber(self, file_data: bytes) -> Tuple[str, List[Dict], List[Dict]]:
        """Извлечение с помощью pdfplumber (отлично для таблиц)"""
        import pdfplumber
        
        text_parts = []
        tables = []
        
        with pdfplumber.open(io.BytesIO(file_data)) as pdf:
            for page_num, page in enumerate(pdf.pages):
                # Извлекаем текст
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                
                # Извлекаем таблицы
                page_tables = page.extract_tables()
                for table_index, table in enumerate(page_tables):
                    if table:
                        # Конвертируем таблицу в удобный формат
                        headers = table[0] if table else []
                        rows = table[1:] if len(table) > 1 else []
                        
                        tables.append({
                            "page": page_num + 1,
                            "index": table_index,
                            "headers": headers,
                            "rows": rows,
                            "raw_data": table
                        })
        
        return '\n\n'.join(text_parts), tables, []
    
    async def _extract_with_pdfminer(self, file_data: bytes) -> Tuple[str, List[Dict], List[Dict]]:
        """Извлечение с помощью PDFMiner"""
        from pdfminer.high_level import extract_text
        
        text = extract_text(io.BytesIO(file_data))
        return text, [], []
    
    async def _extract_with_pypdf2(self, file_data: bytes) -> Tuple[str, List[Dict], List[Dict]]:
        """Извлечение с помощью PyPDF2"""
        try:
            import PyPDF2
            reader_class = PyPDF2.PdfReader
        except ImportError:
            import pypdf
            reader_class = pypdf.PdfReader
        
        reader = reader_class(io.BytesIO(file_data))
        text_parts = []
        
        for page in reader.pages:
            page_text = page.extract_text()
            text_parts.append(page_text)
        
        return '\n\n'.join(text_parts), [], []


class DOCXProcessor(DocumentProcessor):
    """Процессор для DOCX документов"""
    
    def __init__(self):
        super().__init__()
        self.supported_formats = [DocumentFormat.DOCX]
        
    async def process(self, file_data: bytes, metadata: DocumentMetadata) -> ExtractionResult:
        """Обрабатывает DOCX документ"""
        start_time = datetime.now()
        
        try:
            import docx
            
            # Открываем документ
            doc = docx.Document(io.BytesIO(file_data))
            
            # Извлекаем текст
            text_parts = []
            tables = []
            
            # Основной текст
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Таблицы
            for table_index, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    headers = table_data[0] if table_data else []
                    rows = table_data[1:] if len(table_data) > 1 else []
                    
                    tables.append({
                        "index": table_index,
                        "headers": headers,
                        "rows": rows,
                        "raw_data": table_data
                    })
            
            text = '\n\n'.join(text_parts)
            
            # Извлекаем метаданные
            core_props = doc.core_properties
            metadata.author = core_props.author
            metadata.title = core_props.title
            metadata.subject = core_props.subject
            metadata.created_date = core_props.created
            metadata.modified_date = core_props.modified
            
            processing_time = (datetime.now() - start_time).total_seconds()
            metadata.processing_time = processing_time
            metadata.character_count = len(text)
            metadata.word_count = DocumentUtils.count_words(text)
            
            return ExtractionResult(
                text=text,
                metadata=metadata,
                tables=tables,
                success=True,
                extraction_method="DOCX_python_docx",
                processing_details={
                    "processing_time": processing_time,
                    "paragraph_count": len(doc.paragraphs),
                    "table_count": len(tables)
                }
            )
            
        except ImportError:
            logger.error("python-docx library not available. Install: pip install python-docx")
            return ExtractionResult(
                text="",
                metadata=metadata,
                success=False,
                extraction_method="DOCX_FAILED",
                processing_details={"error": "python-docx library not available"}
            )
        except Exception as e:
            logger.error(f"DOCX processing failed: {e}")
            return ExtractionResult(
                text="",
                metadata=metadata,
                success=False,
                extraction_method="DOCX_FAILED",
                processing_details={"error": str(e)}
            )


class TextProcessor(DocumentProcessor):
    """Процессор для текстовых файлов"""
    
    def __init__(self):
        super().__init__()
        self.supported_formats = [DocumentFormat.TXT, DocumentFormat.MARKDOWN, DocumentFormat.RTF]
        
    async def process(self, file_data: bytes, metadata: DocumentMetadata) -> ExtractionResult:
        """Обрабатывает текстовые файлы"""
        start_time = datetime.now()
        
        try:
            # Определяем кодировку
            encoding = DocumentUtils.detect_encoding(file_data)
            
            # Декодируем текст
            text = file_data.decode(encoding, errors='replace')
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            # Обновляем метаданные
            metadata.encoding = encoding
            metadata.processing_time = processing_time
            metadata.character_count = len(text)
            metadata.word_count = DocumentUtils.count_words(text)
            
            return ExtractionResult(
                text=text,
                metadata=metadata,
                success=True,
                extraction_method=f"TEXT_{encoding}",
                processing_details={
                    "encoding": encoding,
                    "processing_time": processing_time,
                    "line_count": text.count('\n') + 1
                }
            )
            
        except Exception as e:
            logger.error(f"Text processing failed: {e}")
            return ExtractionResult(
                text="",
                metadata=metadata,
                success=False,
                extraction_method="TEXT_FAILED",
                processing_details={"error": str(e)}
            )


class CSVProcessor(DocumentProcessor):
    """Процессор для CSV файлов"""
    
    def __init__(self):
        super().__init__()
        self.supported_formats = [DocumentFormat.CSV]
        
    async def process(self, file_data: bytes, metadata: DocumentMetadata) -> ExtractionResult:
        """Обрабатывает CSV файлы"""
        start_time = datetime.now()
        
        try:
            # Определяем кодировку
            encoding = DocumentUtils.detect_encoding(file_data)
            text_data = file_data.decode(encoding, errors='replace')
            
            # Парсим CSV
            csv_reader = csv.reader(io.StringIO(text_data))
            rows = list(csv_reader)
            
            if not rows:
                raise Exception("Empty CSV file")
            
            # Первая строка как заголовки
            headers = rows[0] if rows else []
            data_rows = rows[1:] if len(rows) > 1 else []
            
            # Создаём структурированные данные
            structured_data = {
                "format": "csv",
                "headers": headers,
                "rows": data_rows,
                "row_count": len(data_rows),
                "column_count": len(headers)
            }
            
            # Создаём текстовое представление
            text_parts = [
                f"CSV файл с {len(data_rows)} строками и {len(headers)} колонками",
                f"Заголовки: {', '.join(headers)}",
                "",
                "Данные:"
            ]
            
            # Добавляем первые 10 строк для preview
            for i, row in enumerate(data_rows[:10]):
                text_parts.append(f"Строка {i+1}: {', '.join(str(cell) for cell in row)}")
            
            if len(data_rows) > 10:
                text_parts.append(f"... и ещё {len(data_rows) - 10} строк")
            
            text = '\n'.join(text_parts)
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            # Обновляем метаданные
            metadata.encoding = encoding
            metadata.processing_time = processing_time
            metadata.character_count = len(text_data)
            
            return ExtractionResult(
                text=text,
                metadata=metadata,
                structured_data=structured_data,
                success=True,
                extraction_method="CSV_reader",
                processing_details={
                    "encoding": encoding,
                    "processing_time": processing_time,
                    "row_count": len(data_rows),
                    "column_count": len(headers)
                }
            )
            
        except Exception as e:
            logger.error(f"CSV processing failed: {e}")
            return ExtractionResult(
                text="",
                metadata=metadata,
                success=False,
                extraction_method="CSV_FAILED",
                processing_details={"error": str(e)}
            )


class JSONProcessor(DocumentProcessor):
    """Процессор для JSON файлов"""
    
    def __init__(self):
        super().__init__()
        self.supported_formats = [DocumentFormat.JSON]
        
    async def process(self, file_data: bytes, metadata: DocumentMetadata) -> ExtractionResult:
        """Обрабатывает JSON файлы"""
        start_time = datetime.now()
        
        try:
            # Определяем кодировку
            encoding = DocumentUtils.detect_encoding(file_data)
            text_data = file_data.decode(encoding, errors='replace')
            
            # Парсим JSON
            json_data = json.loads(text_data)
            
            # Создаём читаемое текстовое представление
            text = self._json_to_readable_text(json_data)
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            # Обновляем метаданные
            metadata.encoding = encoding
            metadata.processing_time = processing_time
            metadata.character_count = len(text_data)
            
            return ExtractionResult(
                text=text,
                metadata=metadata,
                structured_data={"format": "json", "data": json_data},
                success=True,
                extraction_method="JSON_parser",
                processing_details={
                    "encoding": encoding,
                    "processing_time": processing_time,
                    "json_type": type(json_data).__name__,
                    "size": len(str(json_data))
                }
            )
            
        except json.JSONDecodeError as e:
            logger.error(f"JSON parsing failed: {e}")
            return ExtractionResult(
                text="",
                metadata=metadata,
                success=False,
                extraction_method="JSON_FAILED",
                processing_details={"error": f"Invalid JSON: {e}"}
            )
        except Exception as e:
            logger.error(f"JSON processing failed: {e}")
            return ExtractionResult(
                text="",
                metadata=metadata,
                success=False,
                extraction_method="JSON_FAILED",
                processing_details={"error": str(e)}
            )
    
    def _json_to_readable_text(self, data, level=0) -> str:
        """Конвертирует JSON в читаемый текст"""
        indent = "  " * level
        lines = []
        
        if isinstance(data, dict):
            lines.append(f"{indent}JSON объект:")
            for key, value in data.items():
                if isinstance(value, (dict, list)):
                    lines.append(f"{indent}  {key}:")
                    lines.append(self._json_to_readable_text(value, level + 2))
                else:
                    lines.append(f"{indent}  {key}: {value}")
        elif isinstance(data, list):
            lines.append(f"{indent}JSON массив ({len(data)} элементов):")
            for i, item in enumerate(data[:5]):  # Показываем первые 5 элементов
                if isinstance(item, (dict, list)):
                    lines.append(f"{indent}  Элемент {i+1}:")
                    lines.append(self._json_to_readable_text(item, level + 2))
                else:
                    lines.append(f"{indent}  Элемент {i+1}: {item}")
            if len(data) > 5:
                lines.append(f"{indent}  ... и ещё {len(data) - 5} элементов")
        else:
            lines.append(f"{indent}{data}")
        
        return '\n'.join(lines) 


class XMLProcessor(DocumentProcessor):
    """Процессор для XML файлов"""
    
    def __init__(self):
        super().__init__()
        self.supported_formats = [DocumentFormat.XML, DocumentFormat.HTML]
        
    async def process(self, file_data: bytes, metadata: DocumentMetadata) -> ExtractionResult:
        """Обрабатывает XML/HTML файлы"""
        start_time = datetime.now()
        
        try:
            # Определяем кодировку
            encoding = DocumentUtils.detect_encoding(file_data)
            text_data = file_data.decode(encoding, errors='replace')
            
            # Пытаемся парсить как XML
            try:
                import xml.etree.ElementTree as ET
                root = ET.fromstring(text_data)
                
                # Извлекаем текст из XML
                text = self._extract_text_from_xml(root)
                
                # Создаём структурированные данные
                structured_data = {
                    "format": "xml",
                    "root_tag": root.tag,
                    "namespace": root.tag.split('}')[0][1:] if '}' in root.tag else None,
                    "element_count": len(list(root.iter())),
                    "attributes": root.attrib
                }
                
            except ET.ParseError:
                # Если не XML, обрабатываем как HTML или просто текст
                if metadata.format == DocumentFormat.HTML:
                    text = self._extract_text_from_html(text_data)
                    structured_data = {"format": "html", "size": len(text_data)}
                else:
                    text = text_data
                    structured_data = {"format": "text", "size": len(text_data)}
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            # Обновляем метаданные
            metadata.encoding = encoding
            metadata.processing_time = processing_time
            metadata.character_count = len(text)
            metadata.word_count = DocumentUtils.count_words(text)
            
            return ExtractionResult(
                text=text,
                metadata=metadata,
                structured_data=structured_data,
                success=True,
                extraction_method="XML_parser",
                processing_details={
                    "encoding": encoding,
                    "processing_time": processing_time,
                    "format": structured_data.get("format", "unknown")
                }
            )
            
        except Exception as e:
            logger.error(f"XML processing failed: {e}")
            return ExtractionResult(
                text="",
                metadata=metadata,
                success=False,
                extraction_method="XML_FAILED",
                processing_details={"error": str(e)}
            )
    
    def _extract_text_from_xml(self, element, level=0) -> str:
        """Извлекает текст из XML элемента рекурсивно"""
        text_parts = []
        indent = "  " * level
        
        # Добавляем тег и атрибуты
        tag_name = element.tag.split('}')[-1] if '}' in element.tag else element.tag
        if element.attrib:
            attrs = ', '.join([f"{k}={v}" for k, v in element.attrib.items()])
            text_parts.append(f"{indent}<{tag_name} {attrs}>")
        else:
            text_parts.append(f"{indent}<{tag_name}>")
        
        # Добавляем текст элемента
        if element.text and element.text.strip():
            text_parts.append(f"{indent}  {element.text.strip()}")
        
        # Обрабатываем дочерние элементы
        for child in element:
            text_parts.append(self._extract_text_from_xml(child, level + 1))
        
        return '\n'.join(text_parts)
    
    def _extract_text_from_html(self, html_data: str) -> str:
        """Извлекает текст из HTML"""
        try:
            # Пытаемся использовать BeautifulSoup если доступно
            try:
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(html_data, 'html.parser')
                
                # Удаляем скрипты и стили
                for script in soup(["script", "style"]):
                    script.decompose()
                
                # Получаем чистый текст
                text = soup.get_text()
                
                # Очищаем лишние пробелы
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = ' '.join(chunk for chunk in chunks if chunk)
                
                return text
                
            except ImportError:
                # Fallback: простое удаление тегов
                import re
                text = re.sub(r'<[^>]+>', '', html_data)
                text = re.sub(r'\s+', ' ', text).strip()
                return text
                
        except Exception as e:
            logger.warning(f"HTML text extraction failed: {e}")
            return html_data


class ImageProcessor(DocumentProcessor):
    """Процессор для изображений с OCR"""
    
    def __init__(self):
        super().__init__()
        self.supported_formats = [DocumentFormat.IMAGE]
        self.ocr_orchestrator = None
        
    async def process(self, file_data: bytes, metadata: DocumentMetadata) -> ExtractionResult:
        """Обрабатывает изображения с помощью OCR"""
        start_time = datetime.now()
        
        try:
            # Инициализируем OCR если ещё не сделали
            if self.ocr_orchestrator is None:
                self.ocr_orchestrator = OCROrchestrator()
            
            # Определяем формат изображения
            image_format = self._detect_image_format(file_data)
            
            # Применяем OCR
            ocr_result = await self.ocr_orchestrator.process_with_best_engine(
                file_data, 
                "rus+eng"
            )
            
            if ocr_result.success:
                processing_time = (datetime.now() - start_time).total_seconds()
                
                # Обновляем метаданные
                metadata.processing_time = processing_time
                metadata.character_count = len(ocr_result.text)
                metadata.word_count = DocumentUtils.count_words(ocr_result.text)
                
                return ExtractionResult(
                    text=ocr_result.text,
                    metadata=metadata,
                    success=True,
                    extraction_method=f"IMAGE_OCR_{ocr_result.engine.value}",
                    processing_details={
                        "image_format": image_format,
                        "ocr_engine": ocr_result.engine.value,
                        "confidence": ocr_result.confidence,
                        "processing_time": processing_time
                    }
                )
            else:
                return ExtractionResult(
                    text="",
                    metadata=metadata,
                    success=False,
                    extraction_method="IMAGE_OCR_FAILED",
                    processing_details={"error": "OCR processing failed"}
                )
                
        except Exception as e:
            logger.error(f"Image processing failed: {e}")
            return ExtractionResult(
                text="",
                metadata=metadata,
                success=False,
                extraction_method="IMAGE_FAILED",
                processing_details={"error": str(e)}
            )
    
    def _detect_image_format(self, file_data: bytes) -> str:
        """Определяет формат изображения"""
        if file_data.startswith(b'\xff\xd8\xff'):
            return "JPEG"
        elif file_data.startswith(b'\x89PNG'):
            return "PNG"
        elif file_data.startswith(b'GIF8'):
            return "GIF"
        elif file_data.startswith(b'BM'):
            return "BMP"
        elif file_data.startswith(b'RIFF') and b'WEBP' in file_data[:12]:
            return "WEBP"
        else:
            return "UNKNOWN"


class DocumentOrchestrator:
    """Оркестратор для координации всех процессоров документов"""
    
    def __init__(self):
        self.processors = {
            DocumentFormat.PDF: PDFProcessor(),
            DocumentFormat.DOCX: DOCXProcessor(),
            DocumentFormat.DOC: DOCXProcessor(),  # Fallback для старых DOC
            DocumentFormat.TXT: TextProcessor(),
            DocumentFormat.MARKDOWN: TextProcessor(),
            DocumentFormat.RTF: TextProcessor(),
            DocumentFormat.CSV: CSVProcessor(),
            DocumentFormat.JSON: JSONProcessor(),
            DocumentFormat.XML: XMLProcessor(),
            DocumentFormat.HTML: XMLProcessor(),  # HTML обрабатывается как XML
            DocumentFormat.IMAGE: ImageProcessor(),
        }
        
        # OCR процессоры для изображений и сканированных документов
        self.ocr_orchestrator = OCROrchestrator()
        
        # Детектор форматов
        self.format_detector = DocumentFormatDetector()
        
        # Кеш для оптимизации
        self._processor_cache = {}
        
    async def process_document(
        self, 
        file_data: Union[bytes, str, Path], 
        filename: Optional[str] = None,
        force_format: Optional[DocumentFormat] = None,
        use_ocr: bool = True,
        ocr_language: str = "rus+eng"
    ) -> ExtractionResult:
        """
        Универсальная обработка документа
        
        Args:
            file_data: Данные файла (bytes, путь к файлу или строка)
            filename: Имя файла для определения формата
            force_format: Принудительный формат (если автодетекция не нужна)
            use_ocr: Использовать OCR для изображений
            ocr_language: Язык для OCR
        """
        start_time = datetime.now()
        
        try:
            # Подготавливаем данные
            if isinstance(file_data, (str, Path)):
                file_path = Path(file_data)
                filename = filename or file_path.name
                file_data = file_path.read_bytes()
            elif isinstance(file_data, str):
                file_data = file_data.encode('utf-8')
            
            # Создаём базовые метаданные
            metadata = DocumentMetadata(
                filename=filename or "unknown",
                file_size=len(file_data),
                file_hash=DocumentUtils.calculate_checksum(file_data)
            )
            
            # Определяем формат документа
            if force_format:
                doc_format = force_format
            else:
                doc_format = self.format_detector.detect_format(file_data, filename)
            
            metadata.format = doc_format
            
            # Проверяем размер файла
            if not DocumentUtils.validate_file_size(file_data):
                return ExtractionResult(
                    text="",
                    metadata=metadata,
                    success=False,
                    extraction_method="SIZE_LIMIT_EXCEEDED",
                    processing_details={"error": "File too large"}
                )
            
            # Выбираем процессор
            processor = self._get_processor(doc_format)
            
            if processor:
                # Обрабатываем документ
                result = await processor.process(file_data, metadata)
                
                # Если результат пустой и это может быть сканированный документ
                if (not result.success or not result.text.strip()) and use_ocr:
                    if doc_format in [DocumentFormat.PDF, DocumentFormat.IMAGE]:
                        logger.info(f"Trying OCR for {doc_format.value}")
                        ocr_result = await self._try_ocr_fallback(file_data, doc_format, ocr_language)
                        if ocr_result and ocr_result.success:
                            # Объединяем результаты
                            result.text = ocr_result.text
                            result.success = True
                            result.extraction_method += "_WITH_OCR"
                            result.processing_details["ocr_used"] = True
                            result.processing_details["ocr_confidence"] = getattr(ocr_result, 'confidence', 0.0)
                
                return result
            else:
                # Неподдерживаемый формат
                return ExtractionResult(
                    text="",
                    metadata=metadata,
                    success=False,
                    extraction_method="UNSUPPORTED_FORMAT",
                    processing_details={"error": f"Format {doc_format.value} not supported"}
                )
                
        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            return ExtractionResult(
                text="",
                metadata=metadata if 'metadata' in locals() else DocumentMetadata(),
                success=False,
                extraction_method="PROCESSING_ERROR",
                processing_details={"error": str(e)}
            )
    
    def _get_processor(self, doc_format: DocumentFormat) -> Optional[DocumentProcessor]:
        """Получает процессор для указанного формата"""
        return self.processors.get(doc_format)
    
    async def _try_ocr_fallback(
        self, 
        file_data: bytes, 
        doc_format: DocumentFormat, 
        language: str
    ) -> Optional[ExtractionResult]:
        """Пытается извлечь текст с помощью OCR"""
        try:
            if doc_format == DocumentFormat.PDF:
                # Конвертируем PDF в изображения и применяем OCR
                ocr_result = await self.ocr_orchestrator.process_with_best_engine(file_data, language)
            elif doc_format == DocumentFormat.IMAGE:
                # Применяем OCR к изображению
                ocr_result = await self.ocr_orchestrator.process_with_best_engine(file_data, language)
            else:
                return None
            
            if ocr_result.success:
                metadata = DocumentMetadata(
                    format=doc_format,
                    processing_time=ocr_result.processing_time,
                    character_count=len(ocr_result.text),
                    word_count=DocumentUtils.count_words(ocr_result.text)
                )
                
                return ExtractionResult(
                    text=ocr_result.text,
                    metadata=metadata,
                    success=True,
                    extraction_method=f"OCR_{ocr_result.engine.value}",
                    processing_details={
                        "ocr_engine": ocr_result.engine.value,
                        "confidence": ocr_result.confidence,
                        "processing_time": ocr_result.processing_time
                    }
                )
            
        except Exception as e:
            logger.warning(f"OCR fallback failed: {e}")
        
        return None
    
    def get_supported_formats(self) -> List[DocumentFormat]:
        """Возвращает список поддерживаемых форматов"""
        return list(self.processors.keys())
    
    def get_processor_info(self) -> Dict[str, Any]:
        """Возвращает информацию о доступных процессорах"""
        info = {}
        for format_type, processor in self.processors.items():
            info[format_type.value] = {
                "class": processor.__class__.__name__,
                "supported_formats": [f.value for f in processor.supported_formats]
            }
        return info


class DocumentTool(BaseTool):
    """
    Универсальный инструмент для работы с документами в KittyCore 3.0
    
    Поддерживает:
    - PDF, DOCX, TXT, RTF, ODT, CSV, JSON, XML
    - OCR для сканированных документов
    - Извлечение таблиц и изображений
    - Метаданные документов
    - Множественные стратегии обработки с fallback
    """
    
    def __init__(self):
        super().__init__()
        self.name = "document_tool"
        self.description = "Универсальный инструмент для работы с документами"
        self.orchestrator = DocumentOrchestrator()
        
        # JSON Schema для инструмента
        self.schema = {
            "type": "object",
            "properties": {
                "action": {
                    "type": "string",
                    "enum": [
                        "extract_text",
                        "extract_metadata", 
                        "extract_tables",
                        "extract_images",
                        "convert_format",
                        "analyze_document",
                        "batch_process",
                        "supported_formats",
                        "get_info"
                    ],
                    "description": "Действие для выполнения"
                },
                "file_path": {
                    "type": "string",
                    "description": "Путь к файлу документа"
                },
                "file_data": {
                    "type": "string",
                    "description": "Base64 закодированные данные файла (альтернатива file_path)"
                },
                "filename": {
                    "type": "string",
                    "description": "Имя файла (для определения формата)"
                },
                "format": {
                    "type": "string",
                    "enum": ["pdf", "docx", "doc", "txt", "markdown", "rtf", "csv", "json", "xml", "image"],
                    "description": "Принудительный формат документа"
                },
                "use_ocr": {
                    "type": "boolean",
                    "description": "Использовать OCR для сканированных документов (по умолчанию true)",
                    "default": True
                },
                "ocr_language": {
                    "type": "string",
                    "description": "Язык для OCR (например: 'rus+eng', 'eng')",
                    "default": "rus+eng"
                },
                "output_format": {
                    "type": "string",
                    "enum": ["text", "json", "markdown"],
                    "description": "Формат вывода результата",
                    "default": "text"
                },
                "include_metadata": {
                    "type": "boolean",
                    "description": "Включить метаданные в результат",
                    "default": True
                },
                "include_tables": {
                    "type": "boolean", 
                    "description": "Включить таблицы в результат",
                    "default": True
                },
                "include_images": {
                    "type": "boolean",
                    "description": "Включить изображения в результат",
                    "default": False
                },
                "batch_files": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "file_path": {"type": "string"},
                            "filename": {"type": "string"}
                        }
                    },
                    "description": "Список файлов для пакетной обработки"
                }
            },
            "required": ["action"],
            "additionalProperties": False
        }
    
    async def execute(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """Выполняет действие с документом"""
        action = params.get("action")
        
        try:
            if action == "supported_formats":
                return await self._get_supported_formats()
            elif action == "get_info":
                return await self._get_info()
            elif action == "batch_process":
                return await self._batch_process(params)
            else:
                # Основные действия с документами
                return await self._process_single_document(action, params)
                
        except Exception as e:
            logger.error(f"DocumentTool execution failed: {e}")
            return {
                "success": False,
                "error": str(e),
                "action": action
            }
    
    async def _process_single_document(self, action: str, params: Dict[str, Any]) -> Dict[str, Any]:
        """Обрабатывает один документ"""
        # Получаем данные файла
        file_data = await self._get_file_data(params)
        if isinstance(file_data, dict) and not file_data.get("success", True):
            return file_data  # Ошибка получения данных
        
        filename = params.get("filename")
        force_format = None
        if params.get("format"):
            try:
                force_format = DocumentFormat(params["format"])
            except ValueError:
                return {
                    "success": False,
                    "error": f"Unsupported format: {params['format']}"
                }
        
        use_ocr = params.get("use_ocr", True)
        ocr_language = params.get("ocr_language", "rus+eng")
        
        # Обрабатываем документ
        result = await self.orchestrator.process_document(
            file_data=file_data,
            filename=filename,
            force_format=force_format,
            use_ocr=use_ocr,
            ocr_language=ocr_language
        )
        
        # Форматируем результат в зависимости от действия
        return self._format_result(action, result, params)
    
    async def _get_file_data(self, params: Dict[str, Any]) -> Union[bytes, Dict[str, Any]]:
        """Получает данные файла из параметров"""
        if "file_path" in params:
            try:
                file_path = Path(params["file_path"])
                if not file_path.exists():
                    return {"success": False, "error": f"File not found: {file_path}"}
                return file_path.read_bytes()
            except Exception as e:
                return {"success": False, "error": f"Failed to read file: {e}"}
        
        elif "file_data" in params:
            try:
                return base64.b64decode(params["file_data"])
            except Exception as e:
                return {"success": False, "error": f"Invalid base64 data: {e}"}
        
        else:
            return {"success": False, "error": "No file_path or file_data provided"}
    
    def _format_result(self, action: str, result: ExtractionResult, params: Dict[str, Any]) -> Dict[str, Any]:
        """Форматирует результат в зависимости от действия"""
        output_format = params.get("output_format", "text")
        include_metadata = params.get("include_metadata", True)
        include_tables = params.get("include_tables", True)
        include_images = params.get("include_images", False)
        
        base_result = {
            "success": result.success,
            "action": action,
            "extraction_method": result.extraction_method
        }
        
        if not result.success:
            base_result["error"] = result.processing_details.get("error", "Unknown error")
            return base_result
        
        # Основной контент в зависимости от действия
        if action == "extract_text":
            base_result["text"] = result.text
            base_result["word_count"] = result.metadata.word_count
            base_result["character_count"] = result.metadata.character_count
        
        elif action == "extract_metadata":
            base_result["metadata"] = {
                "filename": result.metadata.filename,
                "format": result.metadata.format.value if result.metadata.format else None,
                "file_size": result.metadata.file_size,
                "file_hash": result.metadata.file_hash,
                "author": result.metadata.author,
                "title": result.metadata.title,
                "subject": result.metadata.subject,
                "created_date": result.metadata.created_date.isoformat() if result.metadata.created_date else None,
                "modified_date": result.metadata.modified_date.isoformat() if result.metadata.modified_date else None,
                "processing_time": result.metadata.processing_time,
                "character_count": result.metadata.character_count,
                "word_count": result.metadata.word_count,
                "page_count": result.metadata.page_count,
                "encoding": result.metadata.encoding
            }
        
        elif action == "extract_tables":
            base_result["tables"] = result.tables
            base_result["table_count"] = len(result.tables)
        
        elif action == "extract_images":
            base_result["images"] = result.images if include_images else [
                {k: v for k, v in img.items() if k != "data"} 
                for img in result.images
            ]
            base_result["image_count"] = len(result.images)
        
        elif action == "analyze_document":
            # Полный анализ документа
            base_result["text"] = result.text
            base_result["analysis"] = {
                "word_count": result.metadata.word_count,
                "character_count": result.metadata.character_count,
                "language": "auto-detected",  # TODO: добавить определение языка
                "readability_score": None,    # TODO: добавить анализ читаемости
                "table_count": len(result.tables),
                "image_count": len(result.images),
                "processing_time": result.metadata.processing_time
            }
            
            if include_tables and result.tables:
                base_result["tables"] = result.tables
            
            if include_images and result.images:
                base_result["images"] = result.images
        
        # Добавляем метаданные если запрошены
        if include_metadata and action != "extract_metadata":
            base_result["metadata"] = {
                "filename": result.metadata.filename,
                "format": result.metadata.format.value if result.metadata.format else None,
                "file_size": result.metadata.file_size,
                "processing_time": result.metadata.processing_time
            }
        
        # Добавляем детали обработки
        base_result["processing_details"] = result.processing_details
        
        return base_result
    
    async def _batch_process(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """Пакетная обработка документов"""
        batch_files = params.get("batch_files", [])
        if not batch_files:
            return {"success": False, "error": "No batch_files provided"}
        
        results = []
        failed_count = 0
        
        for file_info in batch_files:
            try:
                # Обрабатываем каждый файл
                file_params = {**params, **file_info}
                result = await self._process_single_document("analyze_document", file_params)
                results.append({
                    "file": file_info.get("filename", file_info.get("file_path", "unknown")),
                    "result": result
                })
                
                if not result.get("success", False):
                    failed_count += 1
                    
            except Exception as e:
                results.append({
                    "file": file_info.get("filename", file_info.get("file_path", "unknown")),
                    "result": {"success": False, "error": str(e)}
                })
                failed_count += 1
        
        return {
            "success": True,
            "action": "batch_process",
            "total_files": len(batch_files),
            "successful_files": len(batch_files) - failed_count,
            "failed_files": failed_count,
            "results": results
        }
    
    async def _get_supported_formats(self) -> Dict[str, Any]:
        """Возвращает поддерживаемые форматы"""
        formats = self.orchestrator.get_supported_formats()
        return {
            "success": True,
            "action": "supported_formats",
            "formats": [f.value for f in formats],
            "count": len(formats)
        }
    
    async def _get_info(self) -> Dict[str, Any]:
        """Возвращает информацию об инструменте"""
        return {
            "success": True,
            "action": "get_info",
            "name": self.name,
            "description": self.description,
            "version": "1.0.0",
            "supported_formats": [f.value for f in self.orchestrator.get_supported_formats()],
            "processors": self.orchestrator.get_processor_info(),
            "ocr_engines": [engine.value for engine in OCREngine],
            "capabilities": [
                "text_extraction",
                "metadata_extraction", 
                "table_extraction",
                "image_extraction",
                "ocr_processing",
                "batch_processing",
                "format_detection",
                "encoding_detection"
            ]
        } 


# Экспортируемые классы
__all__ = [
    'DocumentTool',
    'DocumentOrchestrator', 
    'DocumentFormat',
    'DocumentMetadata',
    'ExtractionResult',
    'OCREngine',
    'OCRResult'
]


async def quick_test_document_tool():
    """Быстрый тест DocumentTool с примером"""
    print("🚀 Тестируем DocumentTool...")
    
    tool = DocumentTool()
    
    # Тест 1: Получение информации
    print("\n📋 Тест 1: Получение информации об инструменте")
    result = await tool.execute({"action": "get_info"})
    print(f"✅ Поддерживаемые форматы: {result.get('supported_formats', [])}")
    print(f"✅ Возможности: {len(result.get('capabilities', []))} функций")
    
    # Тест 2: Поддерживаемые форматы
    print("\n📝 Тест 2: Список поддерживаемых форматов")
    result = await tool.execute({"action": "supported_formats"})
    print(f"✅ Форматов поддерживается: {result.get('count', 0)}")
    
    # Тест 3: Обработка простого текста
    print("\n📄 Тест 3: Обработка текстового документа")
    sample_text = "Привет, мир! Это тестовый документ для KittyCore 3.0."
    
    result = await tool.execute({
        "action": "extract_text",
        "file_data": base64.b64encode(sample_text.encode('utf-8')).decode(),
        "filename": "test.txt"
    })
    
    if result.get("success"):
        print(f"✅ Текст извлечён: {result.get('word_count', 0)} слов")
        print(f"✅ Метод: {result.get('extraction_method', 'unknown')}")
    else:
        print(f"❌ Ошибка: {result.get('error', 'unknown')}")
    
    print("\n🎉 Тест DocumentTool завершён!")
    return result


if __name__ == "__main__":
    # Запуск теста при прямом запуске файла
    import asyncio
    asyncio.run(quick_test_document_tool())